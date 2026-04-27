#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate HadirkuGO Whitepaper DOCX with professional formatting."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# STYLE SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif i == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_horizontal_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/></w:pBdr>')
    pPr.append(pBdr)


def add_bullet(doc, text, bold_prefix=None, level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = 'Calibri'
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A3C6E"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF0F7"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_info_box(doc, text):
    """Add a highlighted info box paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.italic = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF0F7"/>')
    p._p.get_or_add_pPr().append(shading)


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WHITE PAPER')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HadirkuGO')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Platform Kehadiran Digital Berbasis QR Code\ndengan Gamifikasi dan AI Assistant')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

add_horizontal_line(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('Versi 1.5  \u2022  April 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dokumen Teknis Sistem')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Daftar Isi', level=1)
add_horizontal_line(doc)

toc_items = [
    ('1.', 'Executive Summary', '3'),
    ('2.', 'Background & Problem Statement', '4'),
    ('3.', 'Objectives / Goals', '5'),
    ('4.', 'System Overview', '7'),
    ('5.', 'Architecture & Technical Design', '10'),
    ('6.', 'Key Features & Innovation', '13'),
    ('7.', 'Implementation / Methodology', '16'),
    ('8.', 'Security & Privacy', '18'),
    ('9.', 'Use Case / Scenario', '19'),
    ('10.', 'Performance & Evaluation', '21'),
    ('11.', 'Roadmap', '22'),
    ('12.', 'Conclusion', '23'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
    run = p.add_run(f'{num} {title}')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    if num in ('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.'):
        run.bold = True

doc.add_page_break()

# ============================================================
# SECTION 1: EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)
add_horizontal_line(doc)

doc.add_paragraph(
    'HadirkuGO adalah platform kehadiran digital berbasis QR Code yang dirancang untuk '
    'mentransformasi cara organisasi, institusi pendidikan, dan bisnis mengelola kehadiran '
    'anggotanya. Sistem ini menggabungkan teknologi QR Code dinamis dengan elemen gamifikasi '
    'komprehensif untuk menciptakan pengalaman kehadiran yang tidak hanya akurat, tetapi juga memotivasi.'
)

doc.add_paragraph(
    'Masalah utama yang diselesaikan oleh HadirkuGO adalah rendahnya engagement dan akurasi '
    'dalam sistem kehadiran konvensional. Sistem absensi tradisional \u2014 baik manual maupun '
    'digital \u2014 cenderung monoton, mudah dimanipulasi, dan tidak memberikan insentif bagi '
    'pengguna untuk hadir secara konsisten.'
)

doc.add_paragraph(
    'HadirkuGO menawarkan solusi melalui pendekatan gamifikasi terintegrasi: setiap aktivitas '
    'kehadiran menghasilkan poin, poin menentukan level pengguna, dan performa kehadiran '
    'ditampilkan dalam leaderboard kompetitif. Sistem ini juga dilengkapi dengan fitur Challenge '
    'antar pengguna, Achievement badges, Quiz interaktif, Reward system, serta SaiQu \u2014 '
    'asisten AI berbasis Google Gemini yang membantu pengguna memahami data kehadiran mereka secara real-time.'
)

doc.add_heading('Keunggulan Utama', level=3)
add_bullet(doc, 'QR Code dinamis dengan masa berlaku 10 detik untuk mencegah manipulasi')
add_bullet(doc, 'Sistem gamifikasi multi-layer (Poin, Level, Leaderboard, Achievement, Reward)')
add_bullet(doc, 'AI Agent (SaiQu) terintegrasi berbasis Google Gemini untuk analisis data personal')
add_bullet(doc, 'Multi-role architecture dengan 5 tipe pengguna (Admin, Owner, Lecturer, Student, Parent)')
add_bullet(doc, 'Multi-tenant business model yang mendukung banyak organisasi dalam satu platform')
add_bullet(doc, 'Notifikasi email otomatis untuk setiap aktivitas check-in dan check-out')
add_bullet(doc, 'Scheduled background processing untuk leaderboard, statistik, dan achievement')

doc.add_page_break()

# ============================================================
# SECTION 2: BACKGROUND & PROBLEM STATEMENT
# ============================================================
doc.add_heading('2. Background & Problem Statement', level=1)
add_horizontal_line(doc)

doc.add_heading('2.1 Latar Belakang', level=2)
doc.add_paragraph(
    'Dalam lingkungan pendidikan dan bisnis modern, pencatatan kehadiran merupakan komponen '
    'kritis dalam evaluasi kinerja dan kedisiplinan. Namun, sistem kehadiran konvensional '
    'menghadapi berbagai tantangan fundamental yang menghambat efektivitasnya.'
)
doc.add_paragraph(
    'Perkembangan teknologi mobile dan QR Code membuka peluang baru untuk menciptakan sistem '
    'kehadiran yang lebih aman, efisien, dan engaging. Dengan penetrasi smartphone yang tinggi, '
    'pendekatan berbasis QR Code menjadi solusi yang aksesibel tanpa memerlukan investasi hardware khusus.'
)

doc.add_heading('2.2 Permasalahan Utama', level=2)

doc.add_heading('Rendahnya Akurasi Pencatatan', level=3)
doc.add_paragraph(
    'Sistem absensi manual (tanda tangan, kartu hadir) sangat rentan terhadap manipulasi seperti '
    'titip absen. Bahkan sistem digital berbasis fingerprint atau kartu RFID masih memiliki celah '
    'keamanan dan memerlukan investasi hardware yang signifikan.'
)

doc.add_heading('Kurangnya Motivasi Kehadiran', level=3)
doc.add_paragraph(
    'Sistem kehadiran tradisional bersifat pasif \u2014 hanya mencatat tanpa memberikan feedback '
    'atau insentif. Pengguna tidak memiliki motivasi intrinsik maupun ekstrinsik untuk meningkatkan '
    'konsistensi kehadiran mereka.'
)

doc.add_heading('Keterbatasan Analisis Data', level=3)
doc.add_paragraph(
    'Data kehadiran yang dikumpulkan jarang dimanfaatkan secara optimal. Organisasi kesulitan '
    'mengidentifikasi pola kehadiran, mengukur produktivitas berbasis durasi, atau memberikan '
    'penghargaan kepada anggota yang konsisten.'
)

doc.add_heading('Fragmentasi Manajemen Multi-Lokasi', level=3)
doc.add_paragraph(
    'Organisasi dengan banyak lokasi operasional membutuhkan sistem terpusat yang mampu mengelola '
    'kehadiran di berbagai titik secara simultan, lengkap dengan analisis per lokasi.'
)

doc.add_heading('2.3 Kelemahan Solusi yang Sudah Ada', level=2)
add_bullet(doc, 'Hanya berfokus pada pencatatan tanpa elemen engagement atau gamifikasi')
add_bullet(doc, 'Tidak menyediakan mekanisme kompetisi sehat untuk memotivasi pengguna')
add_bullet(doc, 'Tidak memiliki kemampuan AI untuk analisis data personal secara kontekstual')
add_bullet(doc, 'Terbatas pada satu model bisnis (single-tenant) tanpa dukungan multi-organisasi')
add_bullet(doc, 'Tidak mendukung multi-role dengan granularitas akses yang berbeda per tipe pengguna')

doc.add_page_break()

# ============================================================
# SECTION 3: OBJECTIVES / GOALS
# ============================================================
doc.add_heading('3. Objectives / Goals', level=1)
add_horizontal_line(doc)

doc.add_heading('3.1 Tujuan Sistem', level=2)

add_table(doc,
    ['No', 'Tujuan', 'Deskripsi'],
    [
        ['1', 'Meningkatkan Akurasi Kehadiran',
         'Menggunakan QR Code dinamis dengan token yang expire dalam 10 detik untuk memastikan kehadiran fisik yang valid dan mencegah titip absen.'],
        ['2', 'Memotivasi Konsistensi',
         'Menerapkan gamifikasi komprehensif (poin, level, leaderboard, achievement) untuk mendorong kehadiran yang konsisten dan kompetitif.'],
        ['3', 'Menyediakan Insight Berbasis Data',
         'Memberikan analisis kehadiran mendalam melalui statistik personal, ranking multi-periode, dan AI assistant (SaiQu).'],
        ['4', 'Mendukung Multi-Organisasi',
         'Menyediakan platform multi-tenant yang dapat digunakan oleh berbagai organisasi secara independen dengan manajemen terpisah.'],
        ['5', 'Membangun Komunitas Kompetitif',
         'Menciptakan ekosistem kompetisi sehat melalui challenge 1v1, leaderboard multi-kategori, dan rival comparison system.'],
    ],
    col_widths=[0.4, 1.8, 4.3]
)

doc.add_heading('3.2 Target Pengguna', level=2)

add_table(doc,
    ['Role', 'Deskripsi', 'Kapabilitas Utama'],
    [
        ['Admin', 'Administrator sistem global',
         'Manajemen pengguna, pengaturan role, kontrol akses global'],
        ['Owner', 'Pemilik bisnis/organisasi',
         'Manajemen bisnis, lokasi absensi, staff, quiz, banner, produk reward, Super Quiz'],
        ['Lecturer', 'Pengajar / Supervisor',
         'Manajemen tim, monitoring kehadiran, evaluasi mahasiswa, QR check-in/out, viewboard'],
        ['Student', 'Anggota / Peserta',
         'Check-in/out via QR, akses gamifikasi lengkap, challenge, quiz, redeem produk'],
        ['Parent', 'Orang tua / Wali',
         'Monitoring kehadiran anak melalui dashboard khusus'],
    ],
    col_widths=[1.0, 1.8, 3.7]
)

doc.add_heading('3.3 Use Case Utama', level=2)
add_bullet(doc, 'Institusi pendidikan yang ingin meningkatkan kedisiplinan kehadiran mahasiswa dengan pendekatan gamifikasi')
add_bullet(doc, 'Organisasi bisnis yang membutuhkan tracking kehadiran karyawan di multi-lokasi dengan analisis durasi')
add_bullet(doc, 'Komunitas atau kelompok kerja yang ingin membangun budaya kehadiran kompetitif melalui challenge dan leaderboard')
add_bullet(doc, 'Lembaga pelatihan yang membutuhkan sistem kehadiran dengan elemen engagement dan reward')

doc.add_page_break()

# ============================================================
# SECTION 4: SYSTEM OVERVIEW
# ============================================================
doc.add_heading('4. System Overview', level=1)
add_horizontal_line(doc)

doc.add_heading('4.1 Gambaran Umum', level=2)
doc.add_paragraph(
    'HadirkuGO adalah aplikasi web full-stack yang dibangun di atas framework Laravel 8 (PHP). '
    'Sistem ini mengadopsi arsitektur monolitik dengan pola MVC (Model-View-Controller) dan '
    'dilengkapi dengan 16 scheduled commands untuk pemrosesan data berkala. Platform ini mendukung '
    'autentikasi via Google OAuth 2.0 melalui Laravel Socialite dan menyediakan dashboard yang '
    'berbeda untuk setiap role pengguna.'
)

add_info_box(doc,
    'Sistem terdiri dari 53 Eloquent Models, 16 Artisan Commands, 10 scheduled tasks, '
    '5 Mail classes, dan 67+ database migration files yang mencerminkan kompleksitas dan '
    'kematangan platform.'
)

doc.add_heading('4.2 Fitur Utama', level=2)

doc.add_heading('A. Modul Kehadiran (Attendance)', level=3)
add_bullet(doc, 'Check-in dan Check-out berbasis QR Code dinamis dengan token unik 64 karakter')
add_bullet(doc, 'Token QR dengan masa berlaku 10 detik \u2014 otomatis deactivate setelah digunakan')
add_bullet(doc, 'Tracking durasi per lokasi (duration_at_location) dan total durasi harian (total_daily_duration)')
add_bullet(doc, 'Multi-location support: pengguna dapat berpindah lokasi dalam satu sesi tanpa checkout')
add_bullet(doc, 'Notifikasi email otomatis saat check-in (CheckinMail) dan check-out (CheckoutMail)')
add_bullet(doc, 'Checkout-Checkin mode: proses checkout dan checkin baru dalam satu scan QR')
add_bullet(doc, 'Kalender kehadiran personal dengan detail per tanggal')
add_bullet(doc, 'Attendance history dengan export PDF dan CSV')

doc.add_heading('B. Modul Gamifikasi', level=3)

p = doc.add_paragraph()
run = p.add_run('Sistem Poin')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph(
    'Setiap check-in menghasilkan 1 poin. Saat check-out, pengguna mendapatkan poin sebesar '
    'total durasi sesi dalam menit. Contoh: sesi 3 jam = 180 poin. Poin dikelola melalui dua '
    'mekanisme: UserPoint (riwayat detail) dan UserPointSummary (agregasi total_points dan current_points).'
)

p = doc.add_paragraph()
run = p.add_run('Sistem Level')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph(
    'Pengguna naik level berdasarkan akumulasi total poin. Setiap level memiliki range '
    'minimum_points dan maximum_points. Level diperiksa secara berkala melalui command '
    'checkuser:level yang memproses pengguna dalam chunk 100 untuk efisiensi memori.'
)

p = doc.add_paragraph()
run = p.add_run('Leaderboard Multi-Kategori')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph(
    'Sistem leaderboard diperbarui setiap jam melalui scheduled command leaderboard:update. '
    'Terdapat 3 jenis leaderboard:'
)
add_bullet(doc, '', bold_prefix='User Leaderboard: ', level=0)
p = doc.add_paragraph('Top Points, Top Levels, Top Sessions (daily/weekly/monthly/yearly), Top Duration (daily/weekly/monthly/yearly)')
p.paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Location Leaderboard: ', level=0)
p = doc.add_paragraph('Top Locations berdasarkan jumlah kunjungan dan total durasi per periode')
p.paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Team Leaderboard: ', level=0)
p = doc.add_paragraph('Top Teams berdasarkan total durasi anggota per periode')
p.paragraph_format.left_indent = Inches(0.75)

doc.add_paragraph(
    'Setiap entri leaderboard menyimpan current_rank, previous_rank (untuk tracking pergerakan), '
    'dan title otomatis berdasarkan peringkat:'
)

add_table(doc,
    ['Peringkat', 'Title', 'Keterangan'],
    [
        ['#1', 'Supreme Champion', 'Peringkat tertinggi'],
        ['#2', 'Elite Grandmaster', 'Runner-up'],
        ['#3', 'Grandmaster', 'Posisi ketiga'],
        ['#4 \u2013 #5', 'Master Elite', 'Top 5'],
        ['#6 \u2013 #10', 'Renowned Expert', 'Top 10'],
        ['#11 \u2013 #20', 'Rising Star', 'Top 20'],
        ['#21 \u2013 #50', 'Honored Contender', 'Top 50'],
    ],
    col_widths=[1.2, 2.0, 3.3]
)

p = doc.add_paragraph()
run = p.add_run('Achievement System')
run.bold = True
run.font.size = Pt(11)
doc.add_paragraph(
    'Achievement diberikan secara otomatis melalui scheduled commands:'
)
add_bullet(doc, '', bold_prefix='Daily MP (Morning Person): ')
doc.add_paragraph('Diberikan kepada pengguna pertama yang check-in setiap hari. Diproses setiap jam oleh command achievement:assign-daily-mp.').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Longest Duration: ')
doc.add_paragraph('Diberikan kepada pengguna dengan durasi kehadiran terlama. Diproses harian pada pukul 23:59.').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Adventure Student: ')
doc.add_paragraph('Diberikan kepada pengguna yang mengunjungi banyak lokasi berbeda. Diproses setiap jam.').paragraph_format.left_indent = Inches(0.75)

doc.add_heading('C. Modul Challenge', level=3)
doc.add_paragraph(
    'Pengguna dapat menantang pengguna lain dalam kompetisi 1v1 dengan dua tipe: '
    'Points (siapa yang mengumpulkan lebih banyak poin) atau Duration (siapa yang memiliki '
    'durasi kehadiran lebih lama). Durasi challenge dapat diatur 1\u20137 hari. Sistem mencatat '
    'statistik lengkap: total wins, total losses, win rate, win streak, dan lose streak.'
)

doc.add_heading('D. Modul Quiz & Super Quiz', level=3)
doc.add_paragraph(
    'Quiz standar dibuat per bisnis dengan pertanyaan multiple choice. Super Quiz adalah mode '
    'quiz dengan tingkat kesulitan lebih tinggi: setiap jawaban salah langsung mengakhiri attempt '
    '(eliminasi). Super Quiz memiliki question_limit dan max_score, serta fitur surrender dan '
    'timeout detection.'
)

doc.add_heading('E. Modul Reward & Redeem', level=3)
doc.add_paragraph(
    'Sistem reward menggunakan algoritma probabilitas dengan fairness logic. Jika pengguna '
    'baru saja mendapatkan reward besar (\u2265100 poin) dalam 7 hari terakhir, probabilitas '
    'mendapatkan reward besar berikutnya dikurangi hingga 95% (dikalikan 0.05). Reward hanya '
    'diberikan untuk sesi kehadiran yang memenuhi durasi minimum 6 jam (360 menit).'
)
doc.add_paragraph(
    'Modul Redeem memungkinkan pengguna menukar current_points dengan produk fisik. Produk '
    'memiliki stock_quantity dan points_required. Proses redemption melalui waiting list '
    'yang harus diapprove oleh Owner.'
)

doc.add_heading('F. SaiQu \u2014 AI Assistant', level=3)
doc.add_paragraph(
    'SaiQu adalah chatbot AI terintegrasi yang dibangun menggunakan arsitektur RAG '
    '(Retrieval-Augmented Generation) dengan Google Gemini API. Sistem ini terdiri dari 3 komponen utama:'
)
add_bullet(doc, '', bold_prefix='GeminiService: ')
doc.add_paragraph(
    'Mengelola komunikasi dengan Gemini API. Mendukung model cascade (gemini-2.5-flash \u2192 '
    'gemini-3-flash-preview \u2192 gemini-2.5-flash-lite) dengan auto-retry pada error 429/503.'
).paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='KnowledgeService: ')
doc.add_paragraph(
    'Membangun konteks data real-time berdasarkan keyword matching dari pertanyaan pengguna. '
    'Mencakup 12 topik: attendance, points, level, leaderboard, team, achievement, quiz, '
    'reward, challenge, statistics, features, dan user info.'
).paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='QuestionValidator: ')
doc.add_paragraph(
    'Memfilter pertanyaan agar hanya menjawab topik terkait sistem HadirkuGO. '
    'Pertanyaan di luar scope ditolak dengan respons yang ramah.'
).paragraph_format.left_indent = Inches(0.75)

doc.add_heading('G. Modul Tim & Organisasi', level=3)
doc.add_paragraph(
    'Setiap bisnis (Business) dapat memiliki banyak tim (Team). Tim memiliki struktur hierarki: '
    'Leader, Manager, dan Member. Leader dapat mentransfer kepemimpinan, menambah/menghapus anggota, '
    'dan membubarkan tim. Manager memiliki akses monitoring tanpa kemampuan struktural. '
    'Attendance recap per tim tersedia dengan export PDF dan CSV.'
)

doc.add_heading('H. Rival Comparison System', level=3)
doc.add_paragraph(
    'Pengguna dapat memilih satu rival untuk perbandingan performa langsung di dashboard. '
    'Sistem juga menyediakan Smart Recommendations yang menyarankan rival berdasarkan dua kriteria: '
    'Close Rivals (pengguna dengan selisih poin \u00b110.000) dan Elite Targets (top 3 di leaderboard).'
)

doc.add_page_break()

# ============================================================
# SECTION 5: ARCHITECTURE & TECHNICAL DESIGN
# ============================================================
doc.add_heading('5. Architecture & Technical Design', level=1)
add_horizontal_line(doc)

doc.add_heading('5.1 Arsitektur Sistem', level=2)
doc.add_paragraph(
    'HadirkuGO mengadopsi arsitektur monolitik berbasis Laravel 8 dengan pola MVC. '
    'Keputusan ini diambil untuk menyederhanakan deployment, mengurangi overhead komunikasi '
    'antar-service, dan mempercepat iterasi pengembangan. Meskipun monolitik, sistem ini '
    'terstruktur dengan baik melalui pemisahan concern yang jelas:'
)

add_table(doc,
    ['Layer', 'Komponen', 'Tanggung Jawab'],
    [
        ['Presentation', 'Blade Views, JavaScript', 'Rendering UI, AJAX interactions, QR Code display'],
        ['Controller', '30+ Controllers (5 role-based groups)', 'Request handling, validation, response formatting'],
        ['Service', 'SaiQu Services (Gemini, Knowledge, Validator)', 'Business logic kompleks, AI integration'],
        ['Model', '53 Eloquent Models', 'Data access, relationships, business rules'],
        ['Database', 'MySQL + 67 Migrations', 'Data persistence, schema management'],
        ['Scheduler', '10 Scheduled Commands', 'Background processing (leaderboard, stats, achievements)'],
        ['Mail', '5 Mailable Classes', 'Email notifications (checkin, checkout, MVP, level-up, redemption)'],
    ],
    col_widths=[1.2, 2.5, 2.8]
)

doc.add_heading('5.2 Komponen Utama', level=2)

doc.add_paragraph(
    'Arsitektur sistem dapat divisualisasikan sebagai berikut:'
)

# Architecture diagram as text
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
arch_text = (
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502                    CLIENT LAYER                      \u2502\n'
    '\u2502  [Browser] \u2192 Blade Views + JS + AJAX               \u2502\n'
    '\u2502  [QR Scanner Device] \u2192 Camera + POST API            \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\n'
    '                          \u2502\n'
    '                          \u25bc\n'
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502                  APPLICATION LAYER                    \u2502\n'
    '\u2502  Middleware: Auth, CSRF, Role-Check, CORS, Locale     \u2502\n'
    '\u2502  Controllers: Admin / Owner / Lecturer / Student      \u2502\n'
    '\u2502  Services: GeminiService, KnowledgeService            \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\n'
    '                          \u2502\n'
    '              \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '              \u25bc                           \u25bc\n'
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502     DATA LAYER          \u2502  \u2502   EXTERNAL SERVICES    \u2502\n'
    '\u2502  MySQL Database          \u2502  \u2502  Google OAuth 2.0      \u2502\n'
    '\u2502  53 Models / 67 Tables   \u2502  \u2502  Google Gemini API     \u2502\n'
    '\u2502  File Storage (avatars)  \u2502  \u2502  SMTP Mail Server      \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\n'
)
run = p.add_run(arch_text)
run.font.name = 'Courier New'
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_heading('5.3 Desain Database (High Level)', level=2)
doc.add_paragraph(
    'Database MySQL terdiri dari 67+ tabel yang dikelola melalui Laravel Migrations. '
    'Berikut adalah entitas utama dan relasinya:'
)

add_table(doc,
    ['Entitas', 'Tabel', 'Relasi Utama'],
    [
        ['User', 'users', 'roles (M2M), attendances (1:N), user_points (1:N), user_level (1:1), point_summary (1:1)'],
        ['Business', 'businesses', 'owner (M:1 User), teams (1:N), attendance_locations (1:N), quizzes (1:N)'],
        ['Team', 'teams', 'business (M:1), leader (M:1 User), members (M2M User), managers (M2M User)'],
        ['Attendance', 'attendances', 'user (M:1), attendance_location (M:1), locations (JSON array)'],
        ['AttendanceToken', 'attendance_tokens', 'user (M:1), token (unique 64-char), expires_at, is_active'],
        ['AttendanceLocation', 'attendance_locations', 'business (M:1), unique_id (UUID), latitude/longitude'],
        ['UserPointSummary', 'user_point_summaries', 'user (1:1), total_points, current_points'],
        ['Level', 'levels', 'name, minimum_points, maximum_points, image_url'],
        ['UserLeaderboard', 'user_leaderboards', 'user (M:1), category, score, current_rank, previous_rank, title'],
        ['Challenge', 'challenges', 'challenger (M:1 User), challenged (M:1 User), type, duration_days, status'],
        ['Quiz / SuperQuiz', 'quizzes / super_quizzes', 'business (M:1), questions (1:N), attempts (1:N)'],
        ['Product', 'products', 'owner (M:1 User), business (M:1), stock_quantity, points_required'],
        ['SaiquConversation', 'saiqu_conversations', 'user (M:1), role (user/model), message'],
    ],
    col_widths=[1.3, 1.8, 3.4]
)

doc.add_heading('5.4 API dan Integrasi', level=2)

add_table(doc,
    ['Integrasi', 'Teknologi', 'Fungsi'],
    [
        ['Google OAuth 2.0', 'Laravel Socialite', 'Autentikasi pengguna via akun Google, auto-create user baru dengan role default Student'],
        ['Google Gemini API', 'HTTP Client (Guzzle)', 'AI chatbot SaiQu dengan model cascade dan RAG context building'],
        ['QR Code Generation', 'SimpleSoftwareIO/simple-qrcode', 'Generate QR Code PNG dari token 64-karakter, output base64'],
        ['Email Notification', 'Laravel Mail + SMTP', 'Notifikasi check-in, check-out, MVP, level-up, redemption request'],
        ['PDF Export', 'barryvdh/laravel-dompdf', 'Export attendance recap per tim dalam format PDF'],
    ],
    col_widths=[1.5, 1.8, 3.2]
)

doc.add_heading('5.5 Model AI \u2014 SaiQu Architecture', level=2)
doc.add_paragraph(
    'SaiQu menggunakan arsitektur RAG (Retrieval-Augmented Generation) yang diimplementasikan '
    'melalui KnowledgeService. Alur pemrosesan:'
)

p = doc.add_paragraph()
run = p.add_run(
    '1. User mengirim pertanyaan via AJAX POST /saiqu/chat\n'
    '2. QuestionValidator memfilter pertanyaan (hanya topik HadirkuGO)\n'
    '3. KnowledgeService melakukan keyword matching terhadap 12 topik\n'
    '4. Data real-time di-query dari database berdasarkan topik yang cocok\n'
    '5. Konteks data (max 2000 karakter) digabung dengan pertanyaan user\n'
    '6. GeminiService mengirim payload ke Gemini API dengan system prompt\n'
    '7. Model cascade: gemini-2.5-flash \u2192 gemini-3-flash-preview \u2192 gemini-2.5-flash-lite\n'
    '8. Response disimpan ke saiqu_conversations untuk history\n'
    '9. History di-trim otomatis (max 5 pasang pesan per user)'
)
run.font.name = 'Calibri'
run.font.size = Pt(10)

add_info_box(doc,
    'Konfigurasi AI: temperature=0.3 (low creativity, high accuracy), '
    'max_output_tokens=1024, rate_limit=50/hari dan 10/menit per user.'
)

doc.add_page_break()

# ============================================================
# SECTION 6: KEY FEATURES & INNOVATION
# ============================================================
doc.add_heading('6. Key Features & Innovation', level=1)
add_horizontal_line(doc)

doc.add_heading('6.1 Fitur Unggulan', level=2)

add_table(doc,
    ['Fitur', 'Deskripsi', 'Inovasi'],
    [
        ['Dynamic QR Token', 'Token 64-karakter dengan TTL 10 detik', 'Mencegah screenshot sharing dan titip absen secara efektif'],
        ['Duration-Based Points', 'Poin = total menit kehadiran', 'Mendorong kehadiran berkualitas, bukan sekadar hadir'],
        ['Multi-Location Session', 'Berpindah lokasi tanpa checkout', 'Tracking durasi per lokasi dalam satu sesi kerja'],
        ['AI-Powered Insights', 'SaiQu chatbot dengan RAG', 'Analisis data personal via natural language dalam Bahasa Indonesia'],
        ['Rival System', 'Perbandingan 1v1 dengan smart recommendations', 'Motivasi kompetitif yang personal dan terukur'],
        ['Fairness Reward Algorithm', 'Probabilitas reward dengan anti-abuse logic', 'Mencegah monopoli reward besar oleh satu pengguna'],
        ['Leaderboard Titles', 'Title otomatis berdasarkan ranking', 'Gamifikasi identitas: Supreme Champion, Grandmaster, dll'],
        ['Morning Person Detection', 'Deteksi check-in pertama harian', 'Mendorong kehadiran pagi dengan achievement dan highlight'],
    ],
    col_widths=[1.5, 2.2, 2.8]
)

doc.add_heading('6.2 Inovasi Utama', level=2)

doc.add_heading('Gamifikasi Berlapis (Multi-Layer Gamification)', level=3)
doc.add_paragraph(
    'HadirkuGO tidak hanya menerapkan satu elemen gamifikasi, tetapi mengintegrasikan '
    'multiple layers yang saling memperkuat:'
)
add_bullet(doc, '', bold_prefix='Layer 1 \u2014 Immediate Reward: ')
doc.add_paragraph('Poin langsung saat check-in (+1) dan check-out (+N menit)').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Layer 2 \u2014 Progression: ')
doc.add_paragraph('Level system yang memberikan sense of growth jangka panjang').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Layer 3 \u2014 Competition: ')
doc.add_paragraph('Leaderboard, challenge 1v1, dan rival comparison').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Layer 4 \u2014 Recognition: ')
doc.add_paragraph('Achievement badges, titles (Supreme Champion, dll), dan frame colors').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Layer 5 \u2014 Tangible Reward: ')
doc.add_paragraph('Produk fisik yang dapat ditukar dengan poin (redeem system)').paragraph_format.left_indent = Inches(0.75)

doc.add_heading('RAG-Based AI Assistant', level=3)
doc.add_paragraph(
    'SaiQu bukan sekadar chatbot generik. Sistem ini menggunakan Retrieval-Augmented Generation '
    'yang secara dinamis membangun konteks dari 12 sumber data real-time berdasarkan topik '
    'pertanyaan pengguna. Ini memungkinkan jawaban yang akurat dan personal tanpa memerlukan '
    'fine-tuning model.'
)

doc.add_heading('6.3 Keunggulan Dibanding Sistem Lain', level=2)

add_table(doc,
    ['Aspek', 'Sistem Konvensional', 'HadirkuGO'],
    [
        ['Metode Absensi', 'Fingerprint / Kartu RFID / Manual', 'QR Code dinamis (10 detik TTL)'],
        ['Motivasi', 'Tidak ada', '5-layer gamifikasi terintegrasi'],
        ['Analisis Data', 'Laporan statis', 'AI chatbot + real-time dashboard + statistik otomatis'],
        ['Kompetisi', 'Tidak ada', 'Leaderboard multi-kategori + Challenge 1v1 + Rival system'],
        ['Multi-Lokasi', 'Terbatas', 'Multi-location session tracking dengan durasi per lokasi'],
        ['Reward', 'Tidak ada', 'Probability-based reward + Point redemption'],
        ['Hardware', 'Perlu perangkat khusus', 'Cukup smartphone + browser'],
        ['Multi-Tenant', 'Jarang', 'Built-in multi-business architecture'],
    ],
    col_widths=[1.3, 2.2, 3.0]
)

doc.add_page_break()

# ============================================================
# SECTION 7: IMPLEMENTATION / METHODOLOGY
# ============================================================
doc.add_heading('7. Implementation / Methodology', level=1)
add_horizontal_line(doc)

doc.add_heading('7.1 Teknologi yang Digunakan', level=2)

add_table(doc,
    ['Kategori', 'Teknologi', 'Versi / Detail'],
    [
        ['Backend Framework', 'Laravel', '8.x (PHP 7.3+ / 8.0+)'],
        ['Database', 'MySQL', 'Relational, 67+ migration files'],
        ['Authentication', 'Laravel Socialite', 'Google OAuth 2.0'],
        ['QR Code', 'SimpleSoftwareIO/simple-qrcode', 'v4.2 \u2014 PNG generation'],
        ['AI / LLM', 'Google Gemini API', 'gemini-2.5-flash (primary), gemini-3-flash-preview (fallback)'],
        ['HTTP Client', 'Guzzle', 'v7.9 \u2014 API calls ke Gemini'],
        ['PDF Generation', 'barryvdh/laravel-dompdf', 'v2.2 \u2014 Attendance recap export'],
        ['Email', 'Laravel Mail', 'SMTP \u2014 5 Mailable classes'],
        ['Caching', 'Laravel Cache', 'File driver \u2014 SaiQu context, stats, achievements'],
        ['Task Scheduling', 'Laravel Scheduler', '10 scheduled commands (hourly + daily)'],
        ['Frontend', 'Blade + JavaScript', 'Server-side rendering + AJAX'],
        ['Asset Bundling', 'Laravel Mix (Webpack)', 'CSS/JS compilation'],
        ['Device Detection', 'jenssegers/agent', 'v2.6 \u2014 Mobile/desktop detection'],
        ['Cloud Storage', 'AWS SDK', 'v3.328 \u2014 File storage support'],
    ],
    col_widths=[1.5, 2.3, 2.7]
)

doc.add_heading('7.2 Metode Implementasi', level=2)

doc.add_heading('Role-Based Access Control (RBAC)', level=3)
doc.add_paragraph(
    'Sistem mengimplementasikan RBAC melalui tabel roles dan role_user (many-to-many). '
    'Middleware CheckRole memvalidasi role pengguna pada setiap request ke route yang dilindungi. '
    'Terdapat 5 role: Admin, Owner, Lecturer, Student, dan Parent. Setiap role memiliki prefix '
    'route dan controller group yang terpisah.'
)

doc.add_heading('Token-Based QR Authentication', level=3)
doc.add_paragraph(
    'Proses check-in/check-out menggunakan mekanisme token ephemeral:'
)
add_bullet(doc, 'Student/Lecturer generate QR Code yang berisi token random 64 karakter (Str::random(64))')
add_bullet(doc, 'Token disimpan di tabel attendance_tokens dengan TTL 10 detik dan status is_active=true')
add_bullet(doc, 'QR Scanner di lokasi membaca token dan mengirim POST request ke server')
add_bullet(doc, 'Server memvalidasi: token exists, is_active=true, dan belum expired')
add_bullet(doc, 'Setelah digunakan, token langsung di-deactivate (is_active=false)')
add_bullet(doc, 'Token lama yang expired dihapus secara otomatis saat generate token baru')

doc.add_heading('Scheduled Background Processing', level=3)
doc.add_paragraph(
    'Sistem menjalankan 10 scheduled commands untuk memproses data secara berkala:'
)

add_table(doc,
    ['Command', 'Jadwal', 'Fungsi'],
    [
        ['leaderboard:update', 'Setiap jam', 'Update semua kategori leaderboard (user, location, team) untuk semua periode'],
        ['leaderboard:sync-frames', 'Setiap jam', 'Sinkronisasi frame colors dan titles berdasarkan ranking terbaru'],
        ['user_statistics:update', 'Setiap jam', 'Update statistik personal (avg checkin time, streak, most frequent location)'],
        ['weekly_rankings:update', 'Setiap jam', 'Update ranking mingguan (total points, sessions, hours)'],
        ['achievement:assign-daily-mp', 'Setiap jam', 'Assign achievement Daily MP ke check-in pertama hari ini'],
        ['achievement:assign-longest-duration', 'Harian 23:59', 'Assign achievement durasi terlama hari ini'],
        ['achievement:assign-adventure-student', 'Setiap jam', 'Assign achievement untuk pengguna multi-lokasi'],
        ['attendance:deactivate-old', 'Harian 00:05', 'Deactivate attendance records yang masih aktif dari hari sebelumnya'],
        ['mvp:send-email', 'Harian 23:59', 'Kirim email congratulations ke MVP harian'],
        ['ranking:daily', 'Harian 23:59', 'Simpan snapshot ranking harian untuk historical tracking'],
    ],
    col_widths=[2.2, 1.2, 3.1]
)

doc.add_heading('7.3 Algoritma Utama', level=2)

doc.add_heading('Algoritma Perhitungan Poin', level=3)
p = doc.add_paragraph()
run = p.add_run(
    'Check-in:  +1 poin (fixed)\n'
    'Check-out: +N poin (N = total durasi sesi dalam menit)\n'
    '\n'
    'Contoh sesi 4 jam:\n'
    '  Check-in  = +1 poin\n'
    '  Check-out = +240 poin (4 jam x 60 menit)\n'
    '  Total     = 241 poin per sesi'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('Algoritma Fairness Reward', level=3)
p = doc.add_paragraph()
run = p.add_run(
    '1. Cek apakah user mendapat reward besar (>=100 poin) dalam 7 hari terakhir\n'
    '2. Generate random number 0-100\n'
    '3. Iterasi reward (sorted by probability ascending):\n'
    '   - Jika user punya recent big reward DAN reward.points >= 100:\n'
    '     probability = probability * 0.05  (dikurangi 95%)\n'
    '   - Cumulative += probability\n'
    '   - Jika random <= cumulative: pick reward ini\n'
    '4. Fallback: jika tidak ada yang terpilih atau stok habis,\n'
    '   berikan consolation reward (reward dengan poin terendah > 0)'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('Algoritma Weekly Ranking', level=3)
doc.add_paragraph(
    'Ranking mingguan menggunakan multi-criteria sorting dengan prioritas:'
)
add_bullet(doc, '', bold_prefix='Prioritas 1: ')
doc.add_paragraph('Total hours (durasi kehadiran terbanyak)').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Prioritas 2: ')
doc.add_paragraph('Total sessions (jumlah sesi terbanyak, sebagai tiebreaker)').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Prioritas 3: ')
doc.add_paragraph('Total points (poin terbanyak, sebagai tiebreaker kedua)').paragraph_format.left_indent = Inches(0.75)

doc.add_page_break()

# ============================================================
# SECTION 8: SECURITY & PRIVACY
# ============================================================
doc.add_heading('8. Security & Privacy', level=1)
add_horizontal_line(doc)

doc.add_heading('8.1 Perlindungan Data', level=2)

add_table(doc,
    ['Mekanisme', 'Implementasi', 'Detail'],
    [
        ['Authentication', 'Google OAuth 2.0 + Laravel Auth', 'Tidak menyimpan password Google; placeholder password di-hash dengan bcrypt'],
        ['Authorization', 'Role-Based Access Control (RBAC)', 'Middleware CheckRole memvalidasi role pada setiap request; 5 role terpisah'],
        ['CSRF Protection', 'Laravel VerifyCsrfToken Middleware', 'Semua form POST dilindungi token CSRF otomatis'],
        ['Cookie Encryption', 'EncryptCookies Middleware', 'Semua cookie dienkripsi secara otomatis oleh Laravel'],
        ['Input Validation', 'Laravel Request Validation', 'Validasi input pada setiap controller method (required, max, exists, dll)'],
        ['Database Transactions', 'DB::beginTransaction / commit / rollback', 'Operasi attendance menggunakan transaksi atomik untuk konsistensi data'],
        ['Password Hashing', 'bcrypt', 'Password placeholder di-hash; autentikasi utama via OAuth'],
    ],
    col_widths=[1.5, 2.2, 2.8]
)

doc.add_heading('8.2 Keamanan QR Token', level=2)
add_bullet(doc, 'Token menggunakan Str::random(64) \u2014 64 karakter random yang sangat sulit ditebak')
add_bullet(doc, 'TTL (Time-To-Live) hanya 10 detik \u2014 window of opportunity sangat kecil')
add_bullet(doc, 'Single-use: token langsung di-deactivate setelah digunakan (is_active = false)')
add_bullet(doc, 'Token expired otomatis dihapus saat user generate token baru')
add_bullet(doc, 'Validasi ganda: is_active check DAN expires_at check pada setiap scan')

doc.add_heading('8.3 Keamanan AI (SaiQu)', level=2)
add_bullet(doc, 'Rate limiting: 50 pesan/hari dan 10 pesan/menit per user untuk mencegah abuse')
add_bullet(doc, 'Question validation: hanya menjawab pertanyaan terkait sistem HadirkuGO')
add_bullet(doc, 'Data sensitif dilindungi: SaiQu tidak pernah mengekspos password, token, atau email pribadi')
add_bullet(doc, 'Gemini safety filters: respons yang melanggar safety policy otomatis diblokir')
add_bullet(doc, 'Context trimming: konteks data dibatasi 2000 karakter untuk mencegah data leakage')
add_bullet(doc, 'Conversation history auto-trim: maksimal 5 pasang pesan per user')

doc.add_heading('8.4 Pengelolaan Data Sensitif', level=2)
add_bullet(doc, 'Field password dan remember_token di-hidden dari serialisasi model User')
add_bullet(doc, 'Email verification timestamp di-cast sebagai datetime untuk validasi')
add_bullet(doc, 'Logging komprehensif untuk audit trail (check-in, check-out, role changes, errors)')
add_bullet(doc, 'Email logs disimpan terpisah (email_logs table) untuk tracking delivery status')

doc.add_page_break()

# ============================================================
# SECTION 9: USE CASE / SCENARIO
# ============================================================
doc.add_heading('9. Use Case / Scenario', level=1)
add_horizontal_line(doc)

doc.add_heading('9.1 Skenario: Mahasiswa Check-in di Kampus', level=2)

add_table(doc,
    ['Langkah', 'Aktor', 'Aksi', 'Sistem Response'],
    [
        ['1', 'Student', 'Login via Google OAuth', 'Redirect ke Student Dashboard'],
        ['2', 'Student', 'Klik menu "Check-in"', 'Generate QR Code (token 64-char, TTL 10 detik)'],
        ['3', 'Student', 'Tunjukkan QR ke scanner lokasi', 'Scanner membaca token'],
        ['4', 'System', 'Validasi token (active + not expired)', 'Token valid \u2192 proses check-in'],
        ['5', 'System', 'Catat attendance record', 'Simpan user_id, location_id, checkin_time, is_active=true'],
        ['6', 'System', 'Tambah poin (+1)', 'Update UserPointSummary (total_points + current_points)'],
        ['7', 'System', 'Kirim notifikasi', 'Email check-in + notifikasi in-app ke tim'],
        ['8', 'System', 'Deactivate token', 'Token is_active = false'],
    ],
    col_widths=[0.7, 0.8, 2.2, 2.8]
)

doc.add_heading('9.2 Skenario: Sesi Kerja Multi-Lokasi', level=2)
doc.add_paragraph(
    'Seorang lecturer check-in di Lab A pada pukul 08:00, kemudian pindah ke Ruang Meeting B '
    'pada pukul 10:00, dan check-out pada pukul 12:00.'
)

add_table(doc,
    ['Waktu', 'Aksi', 'Durasi di Lokasi', 'Poin'],
    [
        ['08:00', 'Check-in di Lab A', '-', '+1 (check-in)'],
        ['10:00', 'Check-in di Ruang Meeting B', 'Lab A: 120 menit', '+1 (check-in)'],
        ['12:00', 'Check-out', 'Meeting B: 120 menit', '+240 (total durasi sesi)'],
        ['', 'TOTAL', '240 menit (4 jam)', '242 poin'],
    ],
    col_widths=[0.8, 2.0, 1.8, 1.9]
)

doc.add_paragraph(
    'Sistem secara otomatis menghitung durasi per lokasi saat pengguna berpindah, dan '
    'menjumlahkan total durasi saat check-out untuk perhitungan poin.'
)

doc.add_heading('9.3 Skenario: Challenge Antar Pengguna', level=2)

add_table(doc,
    ['Langkah', 'Aksi', 'Detail'],
    [
        ['1', 'Student A membuat challenge', 'Tipe: Points, Durasi: 3 hari, Lawan: Student B'],
        ['2', 'Challenge dimulai', 'Status: active, started_at dicatat'],
        ['3', 'Selama 3 hari', 'Kedua pengguna mengumpulkan poin melalui kehadiran normal'],
        ['4', 'Challenge berakhir', 'Sistem membandingkan total poin yang dikumpulkan selama periode'],
        ['5', 'Hasil dicatat', 'Winner dan loser dicatat di challenge_results, statistik diupdate'],
    ],
    col_widths=[0.7, 2.0, 3.8]
)

doc.add_heading('9.4 Skenario: Interaksi dengan SaiQu AI', level=2)

p = doc.add_paragraph()
run = p.add_run(
    'User:  "Berapa total poin aku bulan ini?"\n\n'
    'SaiQu: [Proses Internal]\n'
    '  1. QuestionValidator: topik valid (poin \u2192 sistem HadirkuGO)\n'
    '  2. KnowledgeService: keyword "poin" matched\n'
    '     \u2192 Query UserPointSummary untuk user ini\n'
    '     \u2192 Query Top 5 poin tertinggi untuk perbandingan\n'
    '  3. Context: "POINTS DATA: Poin kamu: Total=15420, Current=8350"\n'
    '  4. GeminiService: kirim ke gemini-2.5-flash\n\n'
    'SaiQu: "Total poin kamu sekarang 15.420 poin! Keren banget \U0001f525\n'
    '        Poin yang bisa dipakai (current) ada 8.350. Terus semangat ya!"'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('9.5 Skenario: Owner Mengelola Bisnis', level=2)
doc.add_paragraph(
    'Seorang Owner mendaftarkan bisnis baru, menambahkan lokasi absensi, mengundang Lecturer '
    'sebagai staff, dan membuat quiz untuk anggota:'
)
add_bullet(doc, 'Buat bisnis baru dengan business_unique_id unik')
add_bullet(doc, 'Tambah attendance locations dengan nama, deskripsi, dan koordinat GPS (latitude/longitude)')
add_bullet(doc, 'Undang Lecturer sebagai staff \u2192 Lecturer dapat mengelola tim di bawah bisnis tersebut')
add_bullet(doc, 'Buat Quiz dan Super Quiz dengan pertanyaan multiple choice')
add_bullet(doc, 'Kelola produk reward (nama, stok, poin yang dibutuhkan)')
add_bullet(doc, 'Approve/reject redemption requests dari pengguna')
add_bullet(doc, 'Kelola banner informasi yang tampil di dashboard anggota')

doc.add_page_break()

# ============================================================
# SECTION 10: PERFORMANCE & EVALUATION
# ============================================================
doc.add_heading('10. Performance & Evaluation', level=1)
add_horizontal_line(doc)

doc.add_heading('10.1 Strategi Optimasi Performa', level=2)

add_table(doc,
    ['Aspek', 'Strategi', 'Implementasi'],
    [
        ['Database Query', 'Eager Loading', 'Penggunaan with() untuk mencegah N+1 query problem pada relasi User, Leaderboard, Attendance'],
        ['Memory Management', 'Chunked Processing', 'Command CheckUserLevel memproses user dalam chunk 100 untuk menghemat memori'],
        ['Caching', 'Laravel Cache (File)', 'SaiQu context (achievements, levels, rewards, stats) di-cache 1 jam (3600 detik)'],
        ['Leaderboard', 'Materialized Cache Tables', 'Leaderboard disimpan di tabel terpisah (user_leaderboards, location_leaderboards, team_leaderboards) dan diupdate hourly'],
        ['Concurrency', 'withoutOverlapping()', 'Scheduled commands menggunakan withoutOverlapping() untuk mencegah race condition'],
        ['Database Integrity', 'Transactions', 'Operasi attendance menggunakan DB::beginTransaction() untuk atomicity'],
        ['Token Cleanup', 'Auto-delete expired', 'Token expired dihapus otomatis saat generate token baru, mencegah table bloat'],
        ['Landing Page Stats', 'Cache 24 jam', 'Statistik landing page di-cache 86400 detik untuk mengurangi query pada halaman publik'],
    ],
    col_widths=[1.3, 1.8, 3.4]
)

doc.add_heading('10.2 Skalabilitas', level=2)
doc.add_paragraph(
    'Arsitektur HadirkuGO dirancang dengan mempertimbangkan skalabilitas:'
)
add_bullet(doc, '', bold_prefix='Leaderboard Limit: ')
doc.add_paragraph('Setiap kategori leaderboard dibatasi top 50 entries untuk menjaga performa query').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Pagination: ')
doc.add_paragraph('Data list menggunakan pagination untuk membatasi jumlah record per request').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Selective Context: ')
doc.add_paragraph('SaiQu hanya membangun konteks untuk topik yang relevan (keyword matching), bukan seluruh database').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Scheduled Processing: ')
doc.add_paragraph('Perhitungan berat (leaderboard, statistics) dilakukan di background, bukan real-time per request').paragraph_format.left_indent = Inches(0.75)

doc.add_heading('10.3 Reliability', level=2)
add_bullet(doc, '', bold_prefix='AI Fallback Chain: ')
doc.add_paragraph('SaiQu menggunakan 3 model Gemini secara cascade \u2014 jika model utama gagal, otomatis switch ke fallback').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Graceful Degradation: ')
doc.add_paragraph('Jika SaiQu conversation table belum ada, sistem tetap mengembalikan jawaban AI tanpa menyimpan history').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Email Resilience: ')
doc.add_paragraph('Kegagalan pengiriman email tidak menghentikan proses attendance \u2014 error di-log dan proses berlanjut').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Transaction Safety: ')
doc.add_paragraph('Semua operasi attendance dibungkus dalam database transaction dengan rollback otomatis pada error').paragraph_format.left_indent = Inches(0.75)

doc.add_page_break()

# ============================================================
# SECTION 11: ROADMAP
# ============================================================
doc.add_heading('11. Roadmap', level=1)
add_horizontal_line(doc)

doc.add_heading('11.1 Rencana Pengembangan', level=2)

add_table(doc,
    ['Fase', 'Timeline', 'Fitur / Improvement'],
    [
        ['Fase 1 \u2014 Foundation', 'Nov 2024 \u2013 Jan 2025',
         'Core attendance (QR check-in/out), User management, Team management, Point system, Level system, Leaderboard dasar, Achievement system, Email notifications'],
        ['Fase 2 \u2014 Engagement', 'Jan 2025 \u2013 Feb 2025',
         'Challenge system, Daily checkin tracking, Testimony, Reward system (probability-based), Quiz & Super Quiz, Banner management, Weekly ranking'],
        ['Fase 3 \u2014 Intelligence', 'Apr 2026',
         'SaiQu AI Assistant (Gemini integration), Advanced leaderboard (multi-kategori, multi-periode), Rival comparison system, Feedback system, Leaderboard titles & frames'],
        ['Fase 4 \u2014 Future', 'Q3\u2013Q4 2026',
         'Mobile app (React Native / Flutter), Push notifications, Geofencing untuk validasi lokasi, API publik untuk integrasi third-party, Advanced analytics dashboard, Gamifikasi tim (team challenges)'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

doc.add_heading('11.2 Fitur Masa Depan', level=2)
add_bullet(doc, '', bold_prefix='Mobile Application: ')
doc.add_paragraph('Aplikasi native untuk iOS dan Android dengan QR scanner built-in').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Geofencing: ')
doc.add_paragraph('Validasi lokasi GPS saat check-in untuk memastikan pengguna berada di area yang benar').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Push Notifications: ')
doc.add_paragraph('Notifikasi real-time untuk achievement, challenge results, dan ranking changes').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Public API: ')
doc.add_paragraph('REST API untuk integrasi dengan sistem HR, LMS, atau platform third-party lainnya').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Advanced Analytics: ')
doc.add_paragraph('Dashboard analytics dengan visualisasi trend, prediksi kehadiran, dan anomaly detection').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='Team Challenges: ')
doc.add_paragraph('Kompetisi antar tim (bukan hanya individu) untuk meningkatkan kolaborasi').paragraph_format.left_indent = Inches(0.75)
add_bullet(doc, '', bold_prefix='SaiQu Enhancement: ')
doc.add_paragraph('Proactive insights (SaiQu memberikan saran tanpa diminta), multi-language support').paragraph_format.left_indent = Inches(0.75)

doc.add_page_break()

# ============================================================
# SECTION 12: CONCLUSION
# ============================================================
doc.add_heading('12. Conclusion', level=1)
add_horizontal_line(doc)

doc.add_paragraph(
    'HadirkuGO merepresentasikan evolusi signifikan dalam domain sistem kehadiran digital. '
    'Dengan menggabungkan teknologi QR Code dinamis, gamifikasi multi-layer, dan kecerdasan '
    'buatan berbasis Google Gemini, platform ini berhasil mentransformasi aktivitas kehadiran '
    'yang tradisionalnya monoton menjadi pengalaman yang engaging, kompetitif, dan bermakna.'
)

doc.add_paragraph(
    'Dari sisi teknis, sistem ini dibangun di atas fondasi yang solid: Laravel 8 sebagai '
    'framework backend, MySQL sebagai database relasional dengan 67+ migration files, '
    'dan arsitektur yang terstruktur dengan 53 Eloquent Models dan 16 Artisan Commands. '
    'Mekanisme keamanan QR token dengan TTL 10 detik secara efektif mencegah manipulasi '
    'kehadiran, sementara scheduled background processing memastikan data leaderboard, '
    'statistik, dan achievement selalu up-to-date tanpa membebani request pengguna.'
)

doc.add_paragraph(
    'Inovasi utama HadirkuGO terletak pada pendekatan gamifikasi berlapis yang menciptakan '
    'multiple feedback loops: poin langsung saat check-in/out (immediate reward), level '
    'progression (long-term growth), leaderboard dan challenge (competition), achievement '
    'badges dan titles (recognition), serta produk fisik yang dapat ditukar dengan poin '
    '(tangible reward). Kombinasi ini menciptakan ekosistem motivasi yang komprehensif.'
)

doc.add_paragraph(
    'Integrasi SaiQu sebagai AI assistant berbasis RAG menambahkan dimensi baru dalam '
    'interaksi pengguna dengan data kehadiran mereka. Pengguna dapat bertanya dalam bahasa '
    'natural (Bahasa Indonesia) dan mendapatkan insight personal yang akurat, menjadikan '
    'data kehadiran lebih accessible dan actionable.'
)

doc.add_paragraph(
    'Dengan roadmap pengembangan yang jelas menuju mobile application, geofencing, dan '
    'public API, HadirkuGO memiliki potensi untuk menjadi platform kehadiran digital '
    'terdepan yang tidak hanya mencatat kehadiran, tetapi juga membangun budaya kehadiran '
    'yang positif dan kompetitif di berbagai organisasi.'
)

add_horizontal_line(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('HadirkuGO \u2014 Transforming Attendance into Achievement')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = 'Calibri'
run.bold = True
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versi 1.5  \u2022  April 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'WHITEPAPER_HadirkuGO.docx')
doc.save(output_path)
print(f'Whitepaper saved to: {output_path}')
